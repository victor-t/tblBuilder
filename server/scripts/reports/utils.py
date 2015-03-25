import abc
import docx
import os
from StringIO import StringIO

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT


class TableMaker(object):
    """
    Helper-object to build tables
    """
    def __init__(self, colWidths, styles=None, numHeaders=1, tblStyle=None, firstRowCaption=True):
        """
            Required:
                - colWidths: list of floats

            Optional:
                - styles: dict of default styles for different fields
                     - keys include: ["title", "header", "body", "subheading"]
                - tblStyle: str (default None)
                - firstRowCaption: bool (default True)
                - numHeaders: int (default 1)
        """
        self.colWidths = colWidths
        self.styles = styles or {}
        self.numHeaders = numHeaders
        self.tblStyle = tblStyle
        self.firstRowCaption = firstRowCaption
        self.cells = []
        self.rows = 0
        self.cols = len(self.colWidths)

    def render(self, doc):
        tbl = doc.add_table(rows=self.rows, cols=self.cols, style=self.tblStyle)

        # set column widths
        tbl.autofit = False
        for i, col in enumerate(tbl.columns):
            col.width = docx.shared.Inches(self.colWidths[i])

        # build cells
        for cell in self.cells:
            cell.render(tbl)

        # mark rows as headers to break on pages
        if self.numHeaders>=1:
            for i in xrange(self.numHeaders):
                tblHeader = docx.oxml.parse_xml(r'<w:tblHeader {} />'.format(
                    docx.oxml.ns.nsdecls('w')))
                tbl.rows[i]._tr.get_or_add_trPr().append(tblHeader)

        # apply caption-style to the first cell in first-row
        if self.firstRowCaption:

            cell =  tbl.cell(0, 0)

            cellPr = cell._tc.get_or_add_tcPr()
            cellPr.append(docx.oxml.parse_xml(
                r'<w:tcBorders {} ><w:top w:val="nil"/><w:left w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(
                    docx.oxml.ns.nsdecls('w'))))
            cellPr.append(docx.oxml.parse_xml(
                r'<w:shd {} w:val="clear" w:color="auto" w:fill="auto"/>'.format(
                    docx.oxml.ns.nsdecls('w'))))

            # left-align text using justified (LEFT doesn't work)
            for p in cell.paragraphs:
                if p.style is None:
                    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW

        return tbl

    def get_style(self, style=None, default=None):
        # determine which style to use for a cell; otherwise return None
        if style:
            return style
        else:
            return self.styles.get(default, None)

    def _add_cell(self, cell):
        self.cells.append(cell)
        self.rows = max(self.rows, cell.row+1)
        return cell

    def _get_width(self, col, colspan):
        if colspan:
            return sum(self.colWidths[col:col+colspan])
        else:
            return self.colWidths[col]

    def new_tbl_title(self, text, style=None):
        for i, row in enumerate(self.colWidths):
            if i == 0:
                cell = {
                    "row": 0,
                    "col": i,
                    "colspan": self.cols,
                    "width": sum(self.colWidths),
                    "text": text
                }
                style = self.get_style(style=style, default="title")
                if style:
                    cell["style"] = style
            else:
                cell = {
                    "row": 0,
                    "col": i,
                    "width": self.colWidths[i],
                    "text": ""
                }
            self._add_cell(CellMaker(**cell))

    def new_th(self, row, col, text, colspan=None, rowspan=None, style=None):
        cell = {"row": row, "col": col}
        if rowspan:
            cell["rowspan"] = rowspan
        if colspan:
            cell["colspan"] = colspan
        cell["width"] = self._get_width(col, colspan)
        style = self.get_style(style=style, default="header")
        if style:
            cell["style"] = style
        if style:
            cell["text"] = text
            cell["style"] = style
        else:
            cell["runs"] = [{ "text": text, "bold": True, "italic": False }]
        return self._add_cell(CellMaker(**cell))

    def new_td_txt(self, row, col, text, rowspan=None, colspan=None, style=None):
        cell = {"row": row, "col": col, "text": text}
        if rowspan:
            cell["rowspan"] = rowspan
        if colspan:
            cell["colspan"] = colspan
        cell["width"] = self._get_width(col, colspan)
        style = self.get_style(style=style, default="body")
        if style:
            cell["style"] = style
        return self._add_cell(CellMaker(**cell))

    def new_td_run(self, row, col, runs, rowspan=None, colspan=None, style=None):
        cell = {"row": row, "col": col, "runs": runs}
        if rowspan:
            cell["rowspan"] = rowspan
        if colspan:
            cell["colspan"] = colspan
        cell["width"] = self._get_width(col, colspan)
        style = self.get_style(style=style, default="body")
        if style:
            cell["style"] = style
        return self._add_cell(CellMaker(**cell))

    @classmethod
    def new_run(cls, txt, newline=True, style=None, b=False, i=False):
        if newline: txt += u"\n"
        return {"text": txt, "style": style, "bold": b, "italic": i}


class CellMaker(object):
    """
    Helper-object to build table-cells
    """

    def __init__(self, **kw):
        """
            Required:
                - row: int
                - col: int
                - width: float

            May have the following fields:
                - rowspan: int
                - colspan: int
                - style: str
                - shade: str
                - text: str
                - runs: list of dictionaries

        """
        self.__dict__.update(kw)

    def render(self, tbl):
        cellD = tbl.cell(self.row, self.col)
        p = cellD.paragraphs[0]

        # merge cells if needed
        rowspan = getattr(self, "rowspan", None)
        colspan = getattr(self, "colspan", None)
        if rowspan or colspan:
            rowIdx = self.row + (rowspan or 1) - 1
            colIdx = self.col + (colspan or 1) - 1
            cellD.merge(tbl.cell(rowIdx, colIdx))

        cellD.width = docx.shared.Inches(self.width)

        # add style
        style = getattr(self, "style", None)
        if style:
            p.style = style

        # add shading
        shade = getattr(self, "shade", None)
        if shade:
            shade_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
                docx.oxml.ns.nsdecls('w'), shade))
            cellD._tc.get_or_add_tcPr().append(shade_elm)

        # add content
        text = getattr(self, "text", None)
        if text:
            p.text = text

        runs = getattr(self, "runs", None)
        if runs:
            for runD in runs:
                run = p.add_run(runD["text"])
                run.bold = runD.get("bold", False)
                run.italic = runD.get("italic", False)
                style = runD.get("style", None)
                if style:
                    run.style = style


class DOCXReport(object):

    def __init__(self, root_path, context):
        self.root_path = root_path
        self.context = context

    def build_report(self):
        """
        Build DOCX report, create content, return file in StringIO format
        """
        fn = os.path.join(self.root_path, self.get_template_fn())
        self.doc = Document(fn)
        self.create_content()

        docx = StringIO()
        self.doc.save(docx)
        docx.seek(0)

        return docx

    @abc.abstractmethod
    def get_template_fn(self):
        """
        Get Word-template filename; template should contain all Word-styles
        used in the report.
        """
        pass

    @abc.abstractmethod
    def create_content(self):
        """
        Main-method called to generate the content in a Word report
        """
        pass

    def setLandscape(self):
        section = self.doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width  = Inches(11)
        section.page_height  = Inches(8.5)
